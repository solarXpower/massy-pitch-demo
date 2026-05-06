"""BiziBid Marketing Letter Generator
Generates Word .docx letters for all BiziBid merchant and bank partners.
Run:  python generate_letters.py
Outputs: one .docx per client in their respective subfolder.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Colour palette ──────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x0B, 0x1F, 0x3A)
BLUE  = RGBColor(0x1A, 0x4A, 0x8A)
GOLD  = RGBColor(0xC9, 0xA8, 0x4C)
GREY  = RGBColor(0x6B, 0x72, 0x80)
BLACK = RGBColor(0x2C, 0x3E, 0x50)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x16, 0xA3, 0x4A)
RED   = RGBColor(0x8B, 0x1A, 0x1A)

BASE = r'C:\AI Software\WebSites'

# ── Helpers ─────────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_run(para, text, bold=False, italic=False,
            size=10, color=BLACK, font='Calibri'):
    run = para.add_run(text)
    run.bold            = bold
    run.italic          = italic
    run.font.name       = font
    run.font.size       = Pt(size)
    run.font.color.rgb  = color
    return run

def add_heading(doc, text, color=NAVY, size=11, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'C9A84C')
    pBdr.append(bot)
    pPr.append(pBdr)
    run = p.add_run(text.upper())
    run.bold           = True
    run.font.name      = 'Calibri'
    run.font.size      = Pt(size)
    run.font.color.rgb = color
    return p

def no_borders(tbl):
    """Remove all visible borders from a table."""
    for row in tbl.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBdr = OxmlElement('w:tcBdr')
            for side in ('top','left','bottom','right','insideH','insideV'):
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'),   'none')
                el.set(qn('w:sz'),    '0')
                el.set(qn('w:space'), '0')
                el.set(qn('w:color'), 'auto')
                tcBdr.append(el)
            tcPr.append(tcBdr)

def gold_left_border(cell):
    """Apply a gold left border to a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBdr = OxmlElement('w:tcBdr')
    left  = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    '18')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), 'C9A84C')
    tcBdr.append(left)
    for side in ('top', 'right', 'bottom'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'none')
        el.set(qn('w:sz'),    '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tcBdr.append(el)
    tcPr.append(tcBdr)


# ── LETTERHEAD builder (shared) ─────────────────────────────────────────────────
def build_letterhead(doc):
    hdr = doc.add_table(rows=1, cols=2)
    hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr.style = 'Table Grid'

    # Left — brand
    lc = hdr.cell(0, 0)
    set_cell_bg(lc, '0B1F3A')
    lc.width = Inches(4)
    lp = lc.paragraphs[0]
    lp.paragraph_format.space_before = Pt(8)
    lp.paragraph_format.space_after  = Pt(0)
    r1 = lp.add_run('INTER'); r1.bold = True; r1.font.name = 'Calibri'
    r1.font.size = Pt(22); r1.font.color.rgb = WHITE
    r2 = lp.add_run('XDB');   r2.bold = True; r2.font.name = 'Calibri'
    r2.font.size = Pt(22); r2.font.color.rgb = GOLD
    # BiziBid sub-brand line
    bp = lc.add_paragraph()
    bp.paragraph_format.space_before = Pt(2)
    bp.paragraph_format.space_after  = Pt(2)
    rb1 = bp.add_run('Presenting ')
    rb1.font.name = 'Calibri'; rb1.font.size = Pt(9)
    rb1.font.color.rgb = RGBColor(0xAA, 0xBB, 0xCC)
    rb2 = bp.add_run('BiziBid')
    rb2.bold = True; rb2.font.name = 'Calibri'; rb2.font.size = Pt(9)
    rb2.font.color.rgb = GOLD
    rb3 = bp.add_run(' — bizibid.com')
    rb3.font.name = 'Calibri'; rb3.font.size = Pt(9)
    rb3.font.color.rgb = RGBColor(0xAA, 0xBB, 0xCC)
    tag = lc.add_paragraph()
    tag.paragraph_format.space_before = Pt(0)
    tag.paragraph_format.space_after  = Pt(8)
    rt = tag.add_run("Barbados' First Retail Auction & Deals App")
    rt.font.name = 'Calibri'; rt.font.size = Pt(7.5)
    rt.font.color.rgb = RGBColor(0x77, 0x88, 0x99)

    # Right — contact
    rc = hdr.cell(0, 1)
    set_cell_bg(rc, '0B1F3A')
    rc.width = Inches(2.8)
    rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for line in ['contact@bizibid.com', 'Tel: 1(246) 241-3771',
                 'www.bizibid.com', 'www.interxdb.com']:
        rp = rc.add_paragraph(line)
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rp.paragraph_format.space_before = Pt(2)
        rp.paragraph_format.space_after  = Pt(2)
        rn = rp.runs[0]
        rn.font.name = 'Calibri'; rn.font.size = Pt(8.5)
        rn.font.color.rgb = RGBColor(0xAA, 0xBB, 0xCC)

    no_borders(hdr)

    # Gold rule below header
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(0)
    rule.paragraph_format.space_after  = Pt(0)
    pPr  = rule._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '12')
    bot.set(qn('w:space'), '0')
    bot.set(qn('w:color'), 'C9A84C')
    pBdr.append(bot)
    pPr.append(pBdr)


# ── FREE OFFER block (shared) ───────────────────────────────────────────────────
def build_free_offer(doc, store_ref):
    fo = doc.add_table(rows=1, cols=2)
    fo.style = 'Table Grid'
    fo.alignment = WD_TABLE_ALIGNMENT.CENTER

    bc = fo.cell(0, 0)
    set_cell_bg(bc, 'C9A84C')
    bc.width = Inches(1.2)
    bc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    bp1 = bc.paragraphs[0]
    bp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bp1.paragraph_format.space_before = Pt(8)
    bp1.paragraph_format.space_after  = Pt(2)
    add_run(bp1, '60', bold=True, color=NAVY, size=26)
    bp2 = bc.add_paragraph()
    bp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bp2.paragraph_format.space_before = Pt(0)
    bp2.paragraph_format.space_after  = Pt(8)
    add_run(bp2, 'DAYS FREE', bold=True, color=NAVY, size=9)

    tc2 = fo.cell(0, 1)
    set_cell_bg(tc2, '0B1F3A')
    tp1 = tc2.paragraphs[0]
    tp1.paragraph_format.space_before = Pt(8)
    tp1.paragraph_format.space_after  = Pt(4)
    add_run(tp1, 'Zero Cost for Your First 60 Days', bold=True, color=WHITE, size=12)
    tp2 = tc2.add_paragraph()
    tp2.paragraph_format.space_before = Pt(0)
    tp2.paragraph_format.space_after  = Pt(8)
    add_run(tp2, 'Full access to the BiziBid platform at no monthly cost for 60 days. A ',
            color=RGBColor(0xCC, 0xCC, 0xCC), size=10)
    add_run(tp2, '$500.00 deposit', bold=True, color=GOLD, size=10)
    add_run(tp2, f' secures primary ad placement within the BiziBid app for {store_ref} '
                 'from day one. Refundable subject to terms.',
            color=RGBColor(0xCC, 0xCC, 0xCC), size=10)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ── PRICING TABLE (shared) ──────────────────────────────────────────────────────
def build_pricing_table(doc):
    add_heading(doc, 'Platform Pricing — After 60-Day Introductory Period (BBD)')
    pt = doc.add_table(rows=5, cols=3)
    pt.style = 'Table Grid'
    pt.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(['Service', 'Description', 'Rate (BBD)']):
        c = pt.cell(0, i)
        set_cell_bg(c, '0B1F3A')
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(5)
        add_run(p, h, bold=True, color=WHITE, size=10)

    rows_data = [
        ('Base Listing Package',  'Up to 25 items listed per week on the BiziBid app',           '$450.00 / month'),
        ('Additional Items',      'Per additional 10 items above the 25-item base',               '$150.00 / month'),
        ('Deal of the Week',      'Featured deal highlighted to all BiziBid users for 7 days',   '$150.00 / week'),
        ('Deal of the Month',     'Premium featured deal — top placement for the full month',     '$200.00 / month'),
    ]
    for ri, (s, d, r) in enumerate(rows_data):
        bg = 'F7F9FC' if ri % 2 == 0 else 'FFFFFF'
        for ci, txt in enumerate([s, d, r]):
            c = pt.cell(ri + 1, ci)
            set_cell_bg(c, bg)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            bold = (ci == 0 or ci == 2)
            col  = NAVY if bold else BLACK
            add_run(p, txt, bold=bold, color=col, size=10)

    note = doc.add_paragraph(
        'All pricing in Barbados Dollars (BBD). Rates fixed for first 12 months for founding partners.')
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after  = Pt(10)
    for r in note.runs:
        r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY


# ── SIGNATURE block (shared) ────────────────────────────────────────────────────
def build_signature(doc):
    sp0 = doc.add_paragraph()
    sp0.paragraph_format.space_before = Pt(4)
    sp0.paragraph_format.space_after  = Pt(6)
    pPr2  = sp0._p.get_or_add_pPr()
    pBdr2 = OxmlElement('w:pBdr')
    top2  = OxmlElement('w:top')
    top2.set(qn('w:val'),   'single')
    top2.set(qn('w:sz'),    '4')
    top2.set(qn('w:space'), '1')
    top2.set(qn('w:color'), 'D8E2EF')
    pBdr2.append(top2)
    pPr2.append(pBdr2)
    add_run(sp0, 'Yours sincerely,', size=10.5, color=BLACK)

    sn = doc.add_paragraph()
    sn.paragraph_format.space_before = Pt(12)
    sn.paragraph_format.space_after  = Pt(2)
    add_run(sn, 'INTER', bold=True, color=NAVY, size=16)
    add_run(sn, 'XDB',   bold=True, color=GOLD, size=16)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_before = Pt(0)
    sub.paragraph_format.space_after  = Pt(2)
    add_run(sub, 'Presenting the ', color=GREY, size=9)
    add_run(sub, 'BiziBid', bold=True, color=GOLD, size=9)
    add_run(sub, ' Platform', color=GREY, size=9)

    st = doc.add_paragraph()
    st.paragraph_format.space_after = Pt(8)
    add_run(st, "Barbados' First Retail Auction & Deals App", color=GREY, size=9)

    sc = doc.add_paragraph()
    sc.paragraph_format.space_after = Pt(4)
    add_run(sc, 'contact@bizibid.com', color=BLUE, size=10)
    add_run(sc, '   |   Tel 1(246) 241-3771   |   www.bizibid.com', color=GREY, size=10)


# ── FOOTER line (shared) ────────────────────────────────────────────────────────
def build_footer(doc):
    ft = doc.add_paragraph()
    ft.paragraph_format.space_before = Pt(14)
    ft.paragraph_format.space_after  = Pt(4)
    pPr  = ft._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top  = OxmlElement('w:top')
    top.set(qn('w:val'),   'single')
    top.set(qn('w:sz'),    '6')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), 'C9A84C')
    pBdr.append(top)
    pPr.append(pBdr)
    ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(ft, 'Bizi', bold=True, color=GOLD, size=10)
    add_run(ft, 'Bid', bold=True, color=NAVY, size=10)
    add_run(ft, ' by INTERXDB  —  Auction & Get Buzz', color=GREY, size=9)
    add_run(ft, '                    Confidential — For Addressee Only', color=GREY, size=9)


# ══════════════════════════════════════════════════════════════════════════════
# MERCHANT LETTER BUILDER
# ══════════════════════════════════════════════════════════════════════════════
def build_letter(company_name, greeting, pitch_url, pitch_label,
                 store_ref, industry_intro, product_examples,
                 closing_sentence, output_path):

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin   = Cm(2.2)
        sec.right_margin  = Cm(2.2)

    build_letterhead(doc)

    # Date + Recipient
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(8)
    add_run(p, 'May 2026', color=GREY, size=10)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    add_run(p, 'The Management Team', bold=True, color=NAVY)
    doc.add_paragraph(company_name).paragraph_format.space_after = Pt(2)
    doc.add_paragraph('Barbados').paragraph_format.space_after = Pt(14)

    # Subject box
    sub_tbl = doc.add_table(rows=1, cols=1)
    sub_tbl.style = 'Table Grid'
    sc = sub_tbl.cell(0, 0)
    set_cell_bg(sc, 'F7F9FC')
    sp1 = sc.paragraphs[0]
    sp1.paragraph_format.space_before = Pt(4)
    sp1.paragraph_format.space_after  = Pt(2)
    add_run(sp1, 'PARTNERSHIP INVITATION', color=GREY, size=8.5, bold=True)
    sp2 = sc.add_paragraph()
    sp2.paragraph_format.space_before = Pt(0)
    sp2.paragraph_format.space_after  = Pt(6)
    add_run(sp2, "Introducing BiziBid — Barbados' First Retail Auction & Deals App",
            bold=True, color=NAVY, size=13)
    gold_left_border(sc)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Salutation
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    add_run(p, f'Dear {greeting},', bold=True, size=11, color=NAVY)

    def body_para(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p, text, size=10.5, color=BLACK)

    body_para(
        f"On behalf of INTERXDB, we are pleased to extend this exclusive invitation to {company_name} "
        "to become one of the founding retail partners on BiziBid — the island's first dedicated "
        "business-to-consumer mobile auction and deals platform."
    )
    body_para(
        f"{industry_intro} The BiziBid app is now available at www.bizibid.com and on iOS and Android. "
        "INTERXDB owns and operates the complete technology — the platform, the BiziBid mobile app, "
        "all hosting, and all technical services. You simply list products — we handle the rest."
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p, 'The BiziBid platform is designed with one purpose: ', size=10.5, color=BLACK)
    add_run(p, f'to keep {store_ref} as the buzz word in every Barbadian household, '
               '24 hours a day, 7 days a week', bold=True, size=10.5, color=NAVY)
    add_run(p, ' — whether your doors are open or not.', size=10.5, color=BLACK)

    # Products block
    if product_examples:
        add_heading(doc, 'Products Well-Suited for BiziBid Auctions')
        for ex in product_examples:
            pb = doc.add_paragraph(style='List Bullet')
            pb.paragraph_format.space_after  = Pt(3)
            pb.paragraph_format.space_before = Pt(0)
            add_run(pb, ex, size=10, color=BLACK)

    # How it works
    add_heading(doc, 'How the BiziBid Platform Works')
    how_items = [
        ('01  App-Only Marketplace',
         'All auctions, Deals of the Week, and Deals of the Month run exclusively through '
         'the BiziBid mobile app (iOS & Android). Download free at bizibid.com.'),
        ('02  Registered Businesses Only',
         'Only verified retail businesses may list products. No private sellers — '
         'a professional marketplace that protects your brand.'),
        ('03  25-Item Weekly Limit',
         'Up to 25 items per store per week. Controlled supply creates scarcity, '
         'drives urgency, and maximises competitive bidding.'),
        ('04  Free Customer Registration',
         'Customers register on BiziBid at no cost. Stores promote the app in-store '
         'and via social media to grow the bidder audience.'),
    ]
    how_tbl = doc.add_table(rows=2, cols=2)
    how_tbl.style = 'Table Grid'
    how_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    idx = 0
    for row in how_tbl.rows:
        for cell in row.cells:
            if idx >= len(how_items): break
            set_cell_bg(cell, 'F7F9FC')
            title, body = how_items[idx]
            tp = cell.paragraphs[0]
            tp.paragraph_format.space_before = Pt(6)
            tp.paragraph_format.space_after  = Pt(4)
            add_run(tp, title, bold=True, color=NAVY, size=10)
            bp2 = cell.add_paragraph(body)
            bp2.paragraph_format.space_before = Pt(0)
            bp2.paragraph_format.space_after  = Pt(6)
            for r in bp2.runs:
                r.font.name = 'Calibri'; r.font.size = Pt(9.5)
                r.font.color.rgb = GREY
            idx += 1
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Free offer
    add_heading(doc, 'Introductory Offer — No Monthly Cost for 60 Days')
    build_free_offer(doc, store_ref)

    # Pricing
    build_pricing_table(doc)

    # Pitch demo (optional)
    if pitch_url:
        add_heading(doc, 'See the Vision — Your Interactive Demo')
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p, f'We have prepared a personalised before-and-after demonstration of the BiziBid '
                   f'platform experience for {company_name}. Visit the link below:',
                size=10.5, color=BLACK)
        dt = doc.add_table(rows=1, cols=1)
        dt.style = 'Table Grid'
        dc = dt.cell(0, 0)
        set_cell_bg(dc, 'F7F9FC')
        dp = dc.paragraphs[0]
        dp.paragraph_format.space_before = Pt(8)
        dp.paragraph_format.space_after  = Pt(8)
        add_run(dp, f'{pitch_label}  →  ', bold=True, color=NAVY, size=11)
        add_run(dp, pitch_url, bold=False, color=BLUE, size=11)
        gold_left_border(dc)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
    else:
        add_heading(doc, 'Live Platform Demo')
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p, 'Visit www.bizibid.com for the full interactive platform walkthrough and live app demo. '
                   'We are happy to arrange an in-person demonstration at your premises.',
                size=10.5, color=BLACK)

    # Obligations
    add_heading(doc, 'Partner Obligations & Platform Conduct')
    ob_tbl = doc.add_table(rows=1, cols=1)
    ob_tbl.style = 'Table Grid'
    oc = ob_tbl.cell(0, 0)
    set_cell_bg(oc, 'FFF8F8')
    oh = oc.paragraphs[0]
    oh.paragraph_format.space_before = Pt(6)
    oh.paragraph_format.space_after  = Pt(4)
    add_run(oh, 'BiziBid Store Partner Responsibilities', bold=True, color=RED, size=10.5)
    for ob in [
        'Actively promote the BiziBid app in-store (signage, POS) and across all social media and marketing channels.',
        'Ensure all listed products are accurately described, available, and fulfilled upon auction completion.',
        'Honour all winning BiziBid auction bids — failure to fulfil may result in suspension from the platform.',
        'Maintain professional communication with winning customers through the BiziBid platform.',
    ]:
        op = oc.add_paragraph(ob, style='List Bullet')
        op.paragraph_format.space_before = Pt(2)
        op.paragraph_format.space_after  = Pt(2)
        for r in op.runs:
            r.font.name = 'Calibri'; r.font.size = Pt(10)
            r.font.color.rgb = BLACK
    pb2 = oc.add_paragraph()
    pb2.paragraph_format.space_before = Pt(6)
    pb2.paragraph_format.space_after  = Pt(6)
    add_run(pb2, 'User Conduct: ', bold=True, color=RED, size=10)
    add_run(pb2, 'Customers who fail to honour confirmed BiziBid auction wins will be banned. '
                 'Disputes are reviewed with evidence from both parties before action is taken.',
            size=10, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Closing
    cp = doc.add_paragraph()
    cp.paragraph_format.space_after = Pt(8)
    cp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(cp, closing_sentence, size=10.5, color=BLACK)

    cp2 = doc.add_paragraph()
    cp2.paragraph_format.space_after = Pt(16)
    add_run(cp2, 'We welcome the opportunity to schedule a BiziBid presentation at your convenience. '
                 'Contact us at contact@bizibid.com or visit www.bizibid.com.',
            size=10.5, color=BLACK)

    build_signature(doc)
    build_footer(doc)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f'  Saved: {output_path}')


# ══════════════════════════════════════════════════════════════════════════════
# BANK / FINANCIAL INSTITUTION LETTER BUILDER (First Citizens, RBC, etc.)
# ══════════════════════════════════════════════════════════════════════════════
def build_bank_letter(company_name, greeting, bank_ref,
                      closing_sentence, output_path):

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin   = Cm(2.2)
        sec.right_margin  = Cm(2.2)

    build_letterhead(doc)

    # Date + Recipient
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(8)
    add_run(p, 'May 2026', color=GREY, size=10)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    add_run(p, 'Asset Management & Recovery Division', bold=True, color=NAVY)
    doc.add_paragraph(company_name).paragraph_format.space_after = Pt(2)
    doc.add_paragraph('Barbados').paragraph_format.space_after = Pt(14)

    # Subject box
    sub_tbl = doc.add_table(rows=1, cols=1)
    sub_tbl.style = 'Table Grid'
    sc = sub_tbl.cell(0, 0)
    set_cell_bg(sc, 'F7F9FC')
    sp1 = sc.paragraphs[0]
    sp1.paragraph_format.space_before = Pt(4)
    sp1.paragraph_format.space_after  = Pt(2)
    add_run(sp1, 'PARTNERSHIP INVITATION — ASSET DISPOSAL', color=GREY, size=8.5, bold=True)
    sp2 = sc.add_paragraph()
    sp2.paragraph_format.space_before = Pt(0)
    sp2.paragraph_format.space_after  = Pt(6)
    add_run(sp2, 'BiziBid — Fast, Transparent Asset Disposal for Financial Institutions',
            bold=True, color=NAVY, size=13)
    gold_left_border(sc)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Salutation
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    add_run(p, f'Dear {greeting},', bold=True, size=11, color=NAVY)

    def body_para(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p, text, size=10.5, color=BLACK)

    body_para(
        f"On behalf of INTERXDB, we are pleased to present BiziBid — Barbados' first mobile auction "
        f"and deals platform — as a highly effective channel for {company_name} to dispose of "
        "repossessed vehicles and foreclosed properties quickly, transparently, and at fair market value."
    )
    body_para(
        "Traditional asset disposal methods — physical auctions, sealed bids, and private negotiations — "
        "are slow, costly, and reach a limited audience. BiziBid solves this by putting your assets in "
        "front of thousands of registered, verified Barbados-based bidders the moment they are listed. "
        "Visit www.bizibid.com to see the platform in action."
    )
    body_para(
        f"Every repossessed vehicle or foreclosed property listed on BiziBid immediately reaches the "
        f"full BiziBid registered user base. Bidding is live, competitive, and fully transparent — "
        f"maximising the recovery value for {bank_ref} on every single asset."
    )

    # Asset types
    add_heading(doc, 'Asset Types Suitable for BiziBid')
    asset_tbl = doc.add_table(rows=1, cols=2)
    asset_tbl.style = 'Table Grid'
    asset_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    assets = [
        ('Repossessed Vehicles',
         'Cars, SUVs, trucks, motorcycles, and commercial vehicles recovered from '
         'defaulted loans. Each unit listed separately with photos, year, make, model, '
         'mileage, and reserve price. Bidding opens immediately.'),
        ('Foreclosed Properties',
         'Residential and commercial properties in foreclosure. Listed as Deal of the Month '
         'for maximum 30-day exposure. Full legal descriptions, photos, and title details '
         'provided by your team and displayed to all registered BiziBid users.'),
    ]
    for i, (title, body) in enumerate(assets):
        c = asset_tbl.cell(0, i)
        set_cell_bg(c, 'F7F9FC')
        tp = c.paragraphs[0]
        tp.paragraph_format.space_before = Pt(6)
        tp.paragraph_format.space_after  = Pt(4)
        add_run(tp, title, bold=True, color=NAVY, size=10)
        bp2 = c.add_paragraph(body)
        bp2.paragraph_format.space_before = Pt(0)
        bp2.paragraph_format.space_after  = Pt(6)
        for r in bp2.runs:
            r.font.name = 'Calibri'; r.font.size = Pt(9.5)
            r.font.color.rgb = GREY
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # How it works for banks
    add_heading(doc, 'How BiziBid Works for Financial Institutions')
    how_items = [
        ('01  Asset Submission',
         'Your team submits asset details (photos, description, reserve price) via the '
         'BiziBid.com portal. INTERXDB reviews and publishes within 24 hours.'),
        ('02  Instant Bidder Reach',
         'The asset immediately appears to all registered BiziBid users in Barbados. '
         'Push notifications alert users watching similar assets.'),
        ('03  Transparent Live Bidding',
         'All bids are live and visible. No closed-door negotiations. Maximum fair '
         'market value is achieved through open competition.'),
        ('04  Fast Disposal',
         'Assets move faster than traditional auctions — reducing holding costs, '
         'administration burden, and legal exposure on stale inventory.'),
    ]
    how_tbl2 = doc.add_table(rows=2, cols=2)
    how_tbl2.style = 'Table Grid'
    how_tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    idx = 0
    for row in how_tbl2.rows:
        for cell in row.cells:
            if idx >= len(how_items): break
            set_cell_bg(cell, 'F7F9FC')
            title, body = how_items[idx]
            tp = cell.paragraphs[0]
            tp.paragraph_format.space_before = Pt(6)
            tp.paragraph_format.space_after  = Pt(4)
            add_run(tp, title, bold=True, color=NAVY, size=10)
            bp2 = cell.add_paragraph(body)
            bp2.paragraph_format.space_before = Pt(0)
            bp2.paragraph_format.space_after  = Pt(6)
            for r in bp2.runs:
                r.font.name = 'Calibri'; r.font.size = Pt(9.5)
                r.font.color.rgb = GREY
            idx += 1
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Free offer
    add_heading(doc, 'Introductory Offer — No Monthly Cost for 60 Days')
    build_free_offer(doc, bank_ref)

    # Pricing (bank version — emphasise Deal of Month for properties)
    add_heading(doc, 'Platform Pricing — After 60-Day Introductory Period (BBD)')
    pt = doc.add_table(rows=5, cols=3)
    pt.style = 'Table Grid'
    pt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['Service', 'Best For', 'Rate (BBD)']):
        c = pt.cell(0, i)
        set_cell_bg(c, '0B1F3A')
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(5)
        add_run(p, h, bold=True, color=WHITE, size=10)
    rows_data = [
        ('Base Listing Package',  'Up to 25 asset listings per week (vehicles)',         '$450.00 / month'),
        ('Additional Listings',   'Per additional 10 assets above the 25-item base',     '$150.00 / month'),
        ('Deal of the Week',      'Featured vehicle — maximum 7-day exposure',           '$150.00 / week'),
        ('Deal of the Month',     'Foreclosed property — full month top placement',      '$200.00 / month'),
    ]
    for ri, (s, d, r) in enumerate(rows_data):
        bg = 'F7F9FC' if ri % 2 == 0 else 'FFFFFF'
        for ci, txt in enumerate([s, d, r]):
            c = pt.cell(ri + 1, ci)
            set_cell_bg(c, bg)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            bold = (ci == 0 or ci == 2)
            add_run(p, txt, bold=bold, color=NAVY if bold else BLACK, size=10)
    note = doc.add_paragraph('All pricing in BBD. Volume pricing available for institutions listing 10+ assets monthly. Contact us to discuss.')
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after  = Pt(10)
    for r in note.runs:
        r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY

    # Demo
    add_heading(doc, 'Live Platform Demo')
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p, 'Visit www.bizibid.com for the full interactive platform walkthrough and live app demo. '
               'We are happy to arrange an in-person demonstration at your offices.',
            size=10.5, color=BLACK)

    # Closing
    cp = doc.add_paragraph()
    cp.paragraph_format.space_after = Pt(8)
    cp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(cp, closing_sentence, size=10.5, color=BLACK)

    cp2 = doc.add_paragraph()
    cp2.paragraph_format.space_after = Pt(16)
    add_run(cp2, 'We welcome the opportunity to schedule a BiziBid presentation at your offices. '
                 'Contact us at contact@bizibid.com or visit www.bizibid.com.',
            size=10.5, color=BLACK)

    build_signature(doc)
    build_footer(doc)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f'  Saved: {output_path}')


# ══════════════════════════════════════════════════════════════════════════════
# GENERATE ALL LETTERS
# ══════════════════════════════════════════════════════════════════════════════
print('Generating BiziBid marketing letters...\n')

# ── 1. GENERAL TEMPLATE ─────────────────────────────────────────────────────
build_letter(
    company_name     = '[Business Name]',
    greeting         = 'Management Team',
    pitch_url        = None,
    pitch_label      = None,
    store_ref        = 'your store',
    industry_intro   = ('Your retail business has built a loyal customer base in Barbados. '
                        'BiziBid gives you a direct mobile channel to keep your customers '
                        'engaged, excited, and bidding — 24 hours a day, 7 days a week.'),
    product_examples = [
        'Electronics — smartphones, laptops, tablets, TVs, smartwatches',
        'Appliances — stoves, microwaves, refrigerators, washing machines',
        'Hardware & Tools — drills, power tools, tool sets, lumber',
        'Furniture & Home — beds, sofas, lamps, mirrors, décor',
        'Food & Beverage — gift baskets, premium beverages, seasonal items',
        'Fashion & Lifestyle — clothing bundles, accessories, gift vouchers',
    ],
    closing_sentence = ('We believe BiziBid represents a significant, low-risk opportunity to extend '
                        'your store\'s digital reach, engage a new generation of customers, and position '
                        'your brand as a leader in Barbados\' evolving retail landscape.'),
    output_path      = fr'{BASE}\Marketing\marketing-letter-general.docx',
)

# ── 2. MASSY STORES ──────────────────────────────────────────────────────────
build_letter(
    company_name     = 'Massy Stores Barbados',
    greeting         = 'Massy Stores Management Team',
    pitch_url        = 'https://interxdb.com/WebSites/massy-pitch/',
    pitch_label      = 'Massy Stores Barbados — Live BiziBid Demo',
    store_ref        = 'Massy Stores',
    industry_intro   = ('As Barbados\' leading supermarket and general retail brand, Massy Stores '
                        'has unmatched reach across the island. BiziBid amplifies that reach by '
                        'placing your products in front of active, engaged bidders every single day.'),
    product_examples = [
        'Electronics & Tech — smartphones, smartwatches, tablets',
        'Home & Living — beds, furniture, mirrors, décor, lamps',
        'Cycling & Outdoors — bicycles, fitness equipment',
        'Food & Beverage — premium gift baskets, rum selections, seasonal bundles',
        'Seasonal & Holiday — Christmas hampers, Easter specials, back-to-school bundles',
    ],
    closing_sentence = ('We believe this partnership represents a significant opportunity to extend '
                        'Massy Stores\' digital reach, engage a new generation of digital-first customers, '
                        'and strengthen your position as Barbados\' leading retail brand on BiziBid.'),
    output_path      = fr'{BASE}\Massy\marketing-letter.docx',
)

# ── 3. CARTER'S HARDWARE ─────────────────────────────────────────────────────
build_letter(
    company_name     = "Carter & Co Ltd",
    greeting         = "Carter & Co Management Team",
    pitch_url        = 'https://interxdb.com/WebSites/carters-pitch/',
    pitch_label      = "Carter & Co Ltd — Live BiziBid Demo",
    store_ref        = "Carter's",
    industry_intro   = ("Carter & Co is one of Barbados' most trusted hardware and homeware destinations. "
                        "BiziBid connects Carter's with thousands of active bidders looking for tools, "
                        "lumber, homeware, and outdoor products at competitive prices."),
    product_examples = [
        'Power Tools — drills, impact drivers, circular saws, tool sets',
        'Hand Tools — hammers, wrenches, toolboxes, hardware accessories',
        'Outdoor & BBQ — grills, outdoor furniture, garden equipment',
        'Home Improvement — paint, fixtures, fittings, electrical supplies',
        'Premium Beverages — wine selections, spirits (where applicable)',
        'Décor & Gifts — crystal ware, vases, decorative items',
    ],
    closing_sentence = ("We believe BiziBid represents a significant opportunity for Carter & Co to extend "
                        "its digital reach, engage a new generation of customers, and build on its strong "
                        "reputation as a trusted Barbados hardware institution."),
    output_path      = fr'{BASE}\Carters\marketing-letter.docx',
)

# ── 4. H&B HARDWARE ──────────────────────────────────────────────────────────
build_letter(
    company_name     = 'ACE H&B Hardware & Lumber Inc.',
    greeting         = 'H&B Hardware Management Team',
    pitch_url        = 'https://interxdb.com/WebSites/hbhardware-pitch/',
    pitch_label      = 'ACE H&B Hardware & Lumber Inc. — Live BiziBid Demo',
    store_ref        = 'H&B Hardware',
    industry_intro   = ('ACE H&B Hardware & Lumber Inc. — "The Helpful Place" — is a trusted name '
                        'for hardware, tools, and lumber across Barbados. BiziBid puts H&B\'s products '
                        'in front of active bidders including contractors, homeowners, and DIY enthusiasts.'),
    product_examples = [
        'Power Tools — DeWalt, Makita, and ACE-brand drills, drivers, and sets',
        'Tool Kits & Cases — professional toolcases, multi-tool combos',
        'Hand Tools & Hardware — all categories of professional hardware',
        'Lumber & Building Materials — selected premium items',
        'Home Improvement Supplies — fixtures, fittings, specialty items',
    ],
    closing_sentence = ("We believe BiziBid represents a significant opportunity for H&B Hardware to extend "
                        "its digital reach, engage a new generation of customers, and reinforce its "
                        'well-earned reputation as "the helpful place" across Barbados.'),
    output_path      = fr'{BASE}\HBHARDWARE\marketing-letter.docx',
)

# ── 5. COURTS BARBADOS LTD ───────────────────────────────────────────────────
build_letter(
    company_name     = 'Courts Barbados Ltd',
    greeting         = 'Courts Barbados Management Team',
    pitch_url        = None,
    pitch_label      = None,
    store_ref        = 'Courts Barbados',
    industry_intro   = ('Courts Barbados is the island\'s leading electronics, appliance, and furniture '
                        'retailer. BiziBid is the perfect channel for Courts to move high-demand products '
                        'through exciting live auctions — generating buzz and urgency around your brand every week.'),
    product_examples = [
        'Smart TVs — 50", 55", 65" QLED, OLED, and LED screens',
        'Large Appliances — stoves, refrigerators, washing machines, dishwashers',
        'Small Appliances — microwaves, air fryers, coffee machines, food warmers',
        'Smartphones & Tablets — Samsung, Apple, and other premium brands',
        'Furniture — bedroom sets, living room suites, dining sets',
        'Laptops & Computing — notebooks, desktops, accessories',
        'Gift Baskets & Bundles — seasonal electronics bundles, tech gift packs',
    ],
    closing_sentence = ('We believe BiziBid represents a significant opportunity for Courts Barbados to '
                        'engage a new generation of digital-first customers, drive weekly excitement '
                        'around your brand, and move high-demand inventory through live competitive auctions.'),
    output_path      = fr'{BASE}\Courts\marketing-letter.docx',
)

# ── 6. PROMOETCH ─────────────────────────────────────────────────────────────
build_letter(
    company_name     = 'PromoTech',
    greeting         = 'PromoTech Management Team',
    pitch_url        = None,
    pitch_label      = None,
    store_ref        = 'PromoTech',
    industry_intro   = ('PromoTech is a specialist technology retailer trusted by Barbados consumers '
                        'for the latest in laptops, smartphones, and accessories. BiziBid creates '
                        'weekly excitement around PromoTech\'s product range — driving sales through '
                        'live competitive bidding on the island\'s premier mobile auction platform.'),
    product_examples = [
        'Laptops — premium ultrabooks, gaming laptops, business notebooks',
        'Smartphones — latest Android and Apple handsets, including bundles',
        'Smartwatches — Apple Watch, Samsung Galaxy Watch, fitness trackers',
        'Tablets & Accessories — iPads, Android tablets, laptop bags, covers',
        'Memory & Storage — SSDs, USB drives, memory cards, external drives',
        'Printers & Peripherals — laser/inkjet printers, scanners, accessories',
    ],
    closing_sentence = ('We believe BiziBid represents a significant opportunity for PromoTech to '
                        'engage Barbados\' growing tech-savvy consumer base, drive weekly demand for '
                        'your product range, and build PromoTech\'s brand as the go-to tech destination '
                        'on Barbados\' premier mobile auction platform.'),
    output_path      = fr'{BASE}\PromoTech\marketing-letter.docx',
)

# ── 7. FIRST CITIZENS BANK ───────────────────────────────────────────────────
build_bank_letter(
    company_name     = 'First Citizens Bank Barbados',
    greeting         = 'First Citizens Bank Asset Management Team',
    bank_ref         = 'First Citizens Bank',
    closing_sentence = ('We believe BiziBid represents a fast, transparent, and cost-effective solution '
                        'for First Citizens Bank to recover maximum value on repossessed vehicles and '
                        'foreclosed properties — reaching more qualified buyers, more quickly, than any '
                        'traditional disposal method available in Barbados today.'),
    output_path      = fr'{BASE}\FirstCitizens\marketing-letter.docx',
)

print('\nAll 7 Word documents generated successfully.')
print('Files saved to their respective client folders.')
